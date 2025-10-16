"""
Utility functions for processing uploaded documents (PDF, Word, TXT)
"""
import io
from typing import Dict, Any
from PyPDF2 import PdfReader
from docx import Document


async def process_uploaded_file(
    file_content: bytes,
    file_name: str,
    file_type: str
) -> Dict[str, Any]:
    """
    Process an uploaded file and extract its content
    
    Args:
        file_content: Raw file content as bytes
        file_name: Name of the file
        file_type: MIME type of the file
        
    Returns:
        Dictionary with extracted content, key points, and citations
    """
    
    extracted_content = ""
    
    try:
        if file_type == "application/pdf":
            extracted_content = await extract_pdf_content(file_content)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            extracted_content = await extract_word_content(file_content)
        elif file_type == "application/msword":
            extracted_content = await extract_word_content(file_content)
        elif file_type == "text/plain":
            extracted_content = file_content.decode('utf-8')
        else:
            extracted_content = f"[Unsupported file type: {file_type}]"
        
        # TODO: Implement key point extraction using NLP
        key_points = []
        
        # TODO: Implement citation extraction
        citations = []
        
        return {
            "content": extracted_content,
            "keyPoints": key_points,
            "citations": citations
        }
        
    except Exception as e:
        print(f"Error processing file {file_name}: {str(e)}")
        return {
            "content": f"[Error extracting content from {file_name}: {str(e)}]",
            "keyPoints": [],
            "citations": []
        }


async def extract_pdf_content(file_content: bytes) -> str:
    """
    Extract text content from a PDF file
    
    Args:
        file_content: PDF file content as bytes
        
    Returns:
        Extracted text content
    """
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PdfReader(pdf_file)
        
        text_content = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_content.append(text)
        
        return "\n\n".join(text_content)
        
    except Exception as e:
        print(f"Error extracting PDF content: {str(e)}")
        raise


async def extract_word_content(file_content: bytes) -> str:
    """
    Extract text content from a Word document (.docx)
    
    Args:
        file_content: Word file content as bytes
        
    Returns:
        Extracted text content
    """
    try:
        doc_file = io.BytesIO(file_content)
        document = Document(doc_file)
        
        text_content = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Also extract text from tables
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n\n".join(text_content)
        
    except Exception as e:
        print(f"Error extracting Word content: {str(e)}")
        raise


async def extract_key_points(text: str, max_points: int = 5) -> list:
    """
    Extract key points from text using NLP
    
    Args:
        text: Text content to analyze
        max_points: Maximum number of key points to extract
        
    Returns:
        List of key points
    """
    # TODO: Implement using spacy or similar NLP library
    # For now, return empty list
    return []


async def extract_citations(text: str) -> list:
    """
    Extract citations from text
    
    Args:
        text: Text content to analyze
        
    Returns:
        List of citations found
    """
    # TODO: Implement citation extraction using regex patterns
    # Look for common citation patterns (APA, MLA, Chicago, etc.)
    return []
